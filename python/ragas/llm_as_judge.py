import requests
import json
import csv
from concurrent.futures import ThreadPoolExecutor
import threading
import re
import time
import subprocess
import atexit
import os
import pandas as pd
from docx.api import Document
from docx.shared import Pt

class LlmJudge:
    def __init__(self):
        self.rag_response = {}
        self.lock = threading.Lock()
        self.rag_answers = {}
        self.llm_answers = {}
        self.limit = 40
        self.file_one = {}
        self.file_two = {}
        self.read_from_files = {}
        self.answers_to_compare = []
        self.evaluator_response = []
        
    def get_queries(self, filename):
        queries = []
        with open(f"data/{filename}.csv", 'r') as file:
            reader = csv.DictReader(file)
            for row in reader:
                query = row.get('Query')
                queries.append(query)
        return queries
    
    def get_limit(self):
        while True:
            try:
                limit = input(f"Enter the limit for the number of queries to process [Default ALL]: ")
                if limit == "":
                    self.limit = 0
                    break
                if limit <= 0:
                    print("Please enter a positive integer.")
                    continue
                else:
                    self.limit = int(limit)
                    break
            except ValueError:
                print("Invalid input. Please enter a positive integer.")
    
    def read_excel_files(self):
        print("Reading Excel files...")
        idx = 0
        for filename in os.listdir("data/answer_quality"):
            if not filename.endswith(".xlsx"):
                continue
            with self.lock:
                df = pd.read_excel(f"data/answer_quality/{filename}")
                for index, row in df.iterrows():
                    query = row.get('Query')
                    answer = row.get('Answers')
                    if idx == 0:
                        self.file_one[query] = answer
                        if len(self.file_one) >= self.limit:
                            idx = 1
                            break
                    elif idx == 1:
                        self.file_two[query] = answer
                        if len(self.file_two) >= self.limit:
                            break
            if len(self.file_one) >= self.limit and len(self.file_two) >= self.limit:
                break
        print(f"Length of file one: {len(self.file_one)}")
        for query, answer in self.file_one.items():
            print(f"Query: {query}")
            if query in self.file_two.keys():
                print("Appending")
                self.answers_to_compare.append(f"""
                        Question: {query}
                        
                        Answer 1: {answer}
                        
                        Answer 2: {self.file_two[query]}
                    """)
        
    def run_evaluator(self):
        with ThreadPoolExecutor(max_workers=1) as executor:
            for key, statement in enumerate(self.answers_to_compare):
                executor.submit(self.compare_answers, key, statement)
            
    
    
    def run_rag(self, queries):
        with ThreadPoolExecutor(max_workers=1) as executor:
            for key, query in enumerate(queries):
                executor.submit(self.process_query, key, query)
                
    def run_judge(self):
        with ThreadPoolExecutor(max_workers=1) as executor:
            for query, answer in self.rag_answers.items():
                if answer['status'] == 'success':
                    executor.submit(self.process_judge_query, query, answer['message'])
                
    def process_query(self, key, query):
        MAX_RETRIES = 3
        INITIAL_RETRY_DELAY = 2  # seconds
        
        try:
            body = self.create_rag_request_body(query)
            url = "http://localhost:8080/predict"
            headers = {'Content-Type': 'application/json'}
            
            retries = 0
            while retries <= MAX_RETRIES:
                try:
                    response = requests.request("POST", url, headers=headers, data=body, timeout=300)
                    response.raise_for_status()
                    break  # exit loop if request is successful
                except (requests.exceptions.RequestException, requests.exceptions.Timeout) as e:
                    retries += 1
                    if retries > MAX_RETRIES:
                        print(f"Request failed for query {key} after {MAX_RETRIES} retries: {e}")
                        return  # exit when all retries are exhausted
                    
                    retry_delay = INITIAL_RETRY_DELAY * (2 ** (retries - 1))
                    print(f"Request attempt {retries} failed for query {key}: {e}. Retrying in {retry_delay} seconds...")
                    time.sleep(retry_delay)
                
            if not response.text:
                print(f"Empty response for query {key}")
                return
                
            data_objects = []
            
            # Split the response text by "data:" and process each part
            parts = response.text.split("data:")
            for part in parts:
                part = part.strip()
                if not part:
                    continue
                    
                # Skip parts that don't look like JSON
                if not (part.startswith('{') or part.startswith('[')):
                    continue
                    
                try:
                    # Parse JSON object from the part
                    obj = json.loads(part)
                    data_objects.append(obj)
                except json.JSONDecodeError as e:
                    # Log problematic content for debugging
                    preview = part[:50] + '...' if len(part) > 50 else part
                    print(f"Failed to parse JSON from part for query {key}: {e}")
                    print(f"Content preview: {preview}")

                    # Optional: Try to fix common JSON issues
                    try:
                        # Sometimes SSE responses have extra text before/after JSON
                        # Try to extract just the JSON portion with regex
                        import re
                        json_match = re.search(r'(\{.*\}|\[.*\])', part, re.DOTALL)
                        if json_match:
                            fixed_json = json_match.group(0)
                            obj = json.loads(fixed_json)
                            data_objects.append(obj)
                            print(f"Successfully extracted valid JSON after initial failure")
                    except:
                        pass  # If this also fails, just skip this part
            
            final_answer = data_objects[-1] if data_objects else None
            if not final_answer:
                print(f"No valid data objects found for query {key}")
                return
                
            message = None
            try:
                message = final_answer["content"]["message"]
            except (KeyError, TypeError):
                print(f"Response missing expected structure for query {key}")
            
            
            with self.lock:
                self.rag_answers[query] = {
                    'status': 'success', # track status in results
                    'message': message
                }
            
        except Exception as e:
            print(f"Unexpected error processing query {key}: {e}")
    
    def process_judge_query(self, query, answer):
        MAX_RETRIES = 3
        INITIAL_RETRY_DELAY = 2  # seconds
        
        try:
            body = json.dumps({
                "query": query,
                "response": answer
            })
            print(f"Sending judge request for query {query} with body: {body}")
            url = "http://localhost:8008/judge"
            headers = {'Content-Type': 'application/json'}
            
            retries = 0
            while retries <= MAX_RETRIES:
                try:
                    response = requests.request("POST", url, headers=headers, data=body, timeout=300)
                    response.raise_for_status()
                    break  # exit loop if request is successful
                except (requests.exceptions.RequestException, requests.exceptions.Timeout) as e:
                    retries += 1
                    if retries > MAX_RETRIES:
                        print(f"Request failed for query {query} after {MAX_RETRIES} retries: {e}")
                        return  # exit when all retries are exhausted
                    
                    retry_delay = INITIAL_RETRY_DELAY * (2 ** (retries - 1))
                    print(f"Request attempt {retries} failed for query {query}: {e}. Retrying in {retry_delay} seconds...")
                    time.sleep(retry_delay)
                
            if not response.text:
                print(f"Empty response for query {query}")
                return
            
            print(f"Response for query {query}: {response.text}")
            
            self.llm_answers[query] = response.text
            
        except Exception as e:
            print(f"Unexpected error processing query {query}: {e}")
            
    def compare_answers(self, query, statement):
        MAX_RETRIES = 3
        INITIAL_RETRY_DELAY = 2  # seconds
        
        try:
            body = json.dumps({
                "statement": statement
            })
            print(f"Sending judge request for query {query} with body: {body}")
            url = "http://localhost:8008/evaluator"
            headers = {'Content-Type': 'application/json'}
            
            retries = 0
            while retries <= MAX_RETRIES:
                try:
                    response = requests.request("POST", url, headers=headers, data=body, timeout=300)
                    response.raise_for_status()
                    break  # exit loop if request is successful
                except (requests.exceptions.RequestException, requests.exceptions.Timeout) as e:
                    retries += 1
                    if retries > MAX_RETRIES:
                        print(f"Request failed for query {query} after {MAX_RETRIES} retries: {e}")
                        return  # exit when all retries are exhausted
                    
                    retry_delay = INITIAL_RETRY_DELAY * (2 ** (retries - 1))
                    print(f"Request attempt {retries} failed for query {query}: {e}. Retrying in {retry_delay} seconds...")
                    time.sleep(retry_delay)
                
            if not response.text:
                print(f"Empty response for query {query}")
                return
            
            self.evaluator_response.append(json.loads(response.text))
            
        except Exception as e:
            print(f"Unexpected error processing query {query}: {e}")
            
    def generate_docx_file(self):
        os.system('cls' if os.name == 'nt' else 'clear')
        # Create a new Document
        document = Document()
        
        # Set the font for the entire document
        style = document.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        timestamp = time.strftime("%Y%m%d-%H%M%S")
        filename = f"data/answer_quality/evaluated_answers_{timestamp}.docx"
        
        print("Generating DOCX file...")
        print(filename)
        
        for index, evaluation in enumerate(self.evaluator_response):
            try:
                print(evaluation)
                # Add the evaluation header
                document.add_paragraph(f"{index + 1}:: {evaluation.get('evaluation')}")
                
                # Add separator
                separator = document.add_paragraph("=" * 65)
                separator.add_run().add_break()  # Add extra line break
                
            except json.JSONDecodeError as e:
                print(f"Failed to decode JSON for evaluation: {evaluation}. Error: {e}")
        
        # Save the document
        document.save(filename)
        print("DOCX file generated successfully.")
        user_input = input("Press Enter to continue...")
            
    def generate_text_file(self):
        os.system('cls' if os.name == 'nt' else 'clear')
        timestamp = time.strftime("%Y%m%d-%H%M%S")
        filename = f"data/answer_quality/answers_{timestamp}.txt"
        for index, evaluation in enumerate(self.evaluator_response):
            try:
                with open(filename, "a") as file:
                    file.write(f"{index + 1}:: {evaluation.get('evaluation')}\n")
                    file.write("=" * 65 + "\n\n")
            except json.JSONDecodeError as e:
                print(f"Failed to decode JSON for evaluation: {evaluation}. Error: {e}")
                
        print("Text file generated successfully.")
        user_input = input("Press Enter to continue...")
    
    @staticmethod       
    def start_fastapi_server():
        print("Starting FastAPI server...")
        process = subprocess.Popen(
            ["uvicorn","llm_judge_api:app","--host","0.0.0.0","--port","8008","--reload"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )
        atexit.register(process.terminate)  # Ensure the process is terminated on exit
        return process

    @staticmethod
    def wait_for_server_ready(url="http://localhost:8008/health", timeout=300):
        start_time = time.time()
        while time.time() - start_time < timeout:
            try:
                response = requests.get(url)
                if response.status_code == 200:
                    print("FastAPI server is ready.")
                    return True
            except requests.exceptions.RequestException:
                pass
            time.sleep(1)
            
        print(f"FastAPI server did not start within {timeout} seconds!")
        return False

    def run_server_in_thread(self):
        server_thread = threading.Thread(target=self.start_fastapi_server)
        server_thread.daemon = True  # Thread must exit when main program exits
        server_thread.start()
        
        if self.wait_for_server_ready(): # We need to wait for the server to be ready before returning
            print("FastAPI server started successfully.")
            return True
        return False
            
    @staticmethod
    def create_rag_request_body(query):
        payload = json.dumps({
            "query": query,
            "normalized_query": "",
            "contextual_query": "",
            "metadata": {},
            "streaming": False,
            "is_first_query": True
        })
        return payload
    
# queries = judge.get_queries("queries")
        # judge.run_rag(queries)
        # judge.run_judge()
        # os.system('cls' if os.name == 'nt' else 'clear')
        # for query, judge_answer in judge.llm_answers.items():
        #     print(f"Query: {query}")
        #     print(f"Judge Answer: {judge_answer}")
        #     print("-" * 40)    
    
if __name__ == "__main__":
    judge = LlmJudge()
    if judge.run_server_in_thread():
        judge.read_excel_files()
        judge.run_evaluator()
        judge.generate_docx_file()
        
    else:
        print("Failed to start LLM Judge server.")